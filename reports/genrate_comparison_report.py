from docx import Document
import os

# ---------------- PATH SETUP ----------------
PROJECT_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
REPORT_PATH = os.path.join(os.path.dirname(__file__), "comparison_report.docx")

# CREATING DOCUMENT
doc = Document()

# Title
doc.add_heading(
    "Legal Document Summarization: Extractive vs Abstractive Comparison",
    level=0
)

# Section 1: Introduction
doc.add_heading("1. Introduction", level=1)
doc.add_paragraph(
    "Legal documents such as contracts, Non-Disclosure Agreements (NDAs), and service "
    "agreements are lengthy and complex. Manual review is time-consuming and prone "
    "to error. This project aims to automate legal document summarization using "
    "Natural Language Processing (NLP) techniques."
)

# Section 2: Dataset
doc.add_heading("2. Dataset Description", level=1)
doc.add_paragraph(
    "The project uses the Contract Understanding Atticus Dataset (CUAD), which "
    "contains real-world legal contracts. Five contracts were selected and the "
    "raw contract text was used for summarization."
)

# Section 3: Methodology
doc.add_heading("3. Methodology", level=1)
doc.add_paragraph(
    "An end-to-end pipeline was developed including data ingestion, preprocessing, "
    "summarization, and evaluation. Long documents were handled using chunking to "
    "overcome transformer input length limitations."
)

# Section 4: Extractive Summarization
doc.add_heading("4. Extractive Summarization", level=1)
doc.add_paragraph(
    "Extractive summarization was implemented using the LexRank algorithm. It selects "
    "important sentences directly from the original document, preserving legal "
    "language and factual accuracy."
)

# Section 5: Abstractive Summarization
doc.add_heading("5. Abstractive Summarization", level=1)
doc.add_paragraph(
    "Abstractive summarization was implemented using the BART transformer model. "
    "This approach generates concise summaries by understanding the semantic meaning "
    "of the document."
)

# Section 6: Evaluation
doc.add_heading("6. Evaluation Strategy", level=1)
doc.add_paragraph(
    "Summaries were evaluated using ROUGE-1, ROUGE-2, and ROUGE-L metrics. These "
    "metrics measure lexical overlap and sequence similarity between extractive "
    "and abstractive summaries."
)

# Section 7: Comparison Table
doc.add_heading("7. Comparative Analysis", level=1)

table = doc.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "Aspect"
hdr_cells[1].text = "Extractive"
hdr_cells[2].text = "Abstractive"

rows = [
    ("Legal Accuracy", "High", "Medium"),
    ("Readability", "Medium", "High"),
    ("Clause Preservation", "Excellent", "Moderate"),
    ("Conciseness", "Low", "High"),
    ("Hallucination Risk", "None", "Possible")
]

for row in rows:
    cells = table.add_row().cells
    cells[0].text = row[0]
    cells[1].text = row[1]
    cells[2].text = row[2]

# Section 8: Conclusion
doc.add_heading("8. Conclusion", level=1)
doc.add_paragraph(
    "The project demonstrates a complete legal document summarization pipeline. "
    "Extractive methods ensure accuracy, while abstractive methods improve readability. "
    "A hybrid approach can significantly improve legal document review efficiency."
)

# Saving document
doc.save(REPORT_PATH)

print("✅ comparison_report.docx generated successfully")

